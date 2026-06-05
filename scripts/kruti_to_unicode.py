#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Kruti Dev 010 -> Unicode Devanagari converter.

Vishnu ji's reference court filings are typed in the legacy **Kruti Dev 010**
font. That font is NOT Unicode: it stores ordinary Latin/ASCII bytes (plus a
few Latin-1 glyphs) that only *look* like Devanagari when rendered in the Kruti
Dev typeface. Reading such a .docx with python-docx therefore yields gibberish
like ``ekuuh; mPp U;k;ky;`` which actually means ``माननीय उच्च न्यायालय``.

This module decodes that legacy encoding to real Unicode so we can read the
filings, extract fields, and build deterministic bilingual specs from them.

It is OFFLINE tooling — it is intentionally NOT imported by ``headnote.api.app``
and cannot affect the live server boot path.

Usage
-----
    python scripts/kruti_to_unicode.py "path/to/file.docx"     # decode a .docx
    python scripts/kruti_to_unicode.py --text "ekuuh; mPp"     # decode a string
    echo "ekuuh;" | python scripts/kruti_to_unicode.py         # decode stdin

Notes on this typist's workflow
-------------------------------
* श (sha) is typed on the apostrophe key and ष (ssha) on the double-quote key,
  the standard Kruti Dev 010 convention. Microsoft Word "smart quotes" then
  rewrites those quotes to curly variants. In this library the curly quotes
  (' ' " ") all resolve to the श-lead and a *straight* double-quote (")
  resolves to the ष-lead — which is what every sample filing actually uses
  (प्रदेश, आदेश via the curly double-quote; अभिभाषक via the straight one).
"""
from __future__ import annotations

import re
import sys

# Ordered replacement table. Order matters: longer / special sequences must be
# listed before the single characters they contain.
PAIRS: list[tuple[str, str]] = [
    # smart quotes used by this typist as the 'sha' lead
    ("”", "'"), ("“", "'"), ("’", "'"), ("‘", "'"),
    # ASCII ligatures that must convert before the bracket remap below
    (")", "द्ध"), (":", "रु"),
    # special conjunct glyphs (Latin-1 / extended)
    ("Ø", "क्र"), ("æ", "क्र"), ("ø", "क्र"), ("=", "त्र"), ("Ý", "द्र"),
    ("™", "त्त"), ("¶", "फ़"),
    # latin-1 brackets / marks
    ("¼", "("), ("½", ")"), ("¡", "ँ"),
    # independent vowel sequences (longest first)
    ("vkS", "औ"), ("vks", "ओ"), ("vk", "आ"), ("bZ", "ई"), (",s", "ऐ"),
    # two-char consonant clusters + o-matra combos
    ("{k", "क्ष"), ("'k", "श"), ('"k', "ष"), ("[k", "ख"), ("Fk", "थ"),
    ("Hk", "भ"), ("?k", "घ"), ("/k", "ध"), (".k", "ण"), ("ks", "ो"),
    ("kS", "ौ"),
    # half / conjunct uppercase consonants
    ("D", "क्"), ("X", "ग्"), ("P", "च्"), ("T", "ज्"), ("F", "थ्"),
    ("U", "न्"), ("I", "प्"), ("C", "ब्"), ("H", "भ्"), ("E", "म्"),
    ("Y", "ल्"), ("O", "व्"), ("L", "स्"), ("R", "त्"), ("'", "श्"),
    ('"', "ष्"), (".", "ण्"), ("[", "ख्"), ("/", "ध्"), ("{", "क्ष्"),
    ("?", "घ्"), ("J", "श्र"),
    # full consonants
    ("d", "क"), ("x", "ग"), ("p", "च"), ("t", "ज"), ("V", "ट"),
    ("B", "ठ"), ("M", "ड"), ("<", "ढ"), ("r", "त"), ("n", "द"),
    ("u", "न"), ("i", "प"), ("Q", "फ"), ("c", "ब"), ("e", "म"),
    (";", "य"), ("j", "र"), ("y", "ल"), ("o", "व"), ("l", "स"),
    ("g", "ह"), ("N", "छ"), ("K", "ज्ञ"), (">", "झ"), ("z", "्र"),
    # independent vowels (single)
    ("v", "अ"), ("b", "इ"), ("m", "उ"), ("Å", "ऊ"), (",", "ए"), ("_", "ऋ"),
    # matras
    ("k", "ा"), ("h", "ी"), ("q", "ु"), ("w", "ू"), ("s", "े"),
    ("S", "ै"), ("`", "ृ"), ("a", "ं"), ("f", "ि"),
    # punctuation / halant / danda
    ("~", "्"), ("A", "।"), ("|", "॥"), ("]", ","), ("@", "/"),
    ("}", "द्व"), ("%", "ः"), ("&", "—"),
]

# Consonant set used by the reordering passes below.
_CONS = "कखगघङचछजझञटठडढणतथदधनपफबभमयरलवशषसहड़ढ़फ़ज़"
_MATRAS = "ािीुूृेैोौंःँ"

# 'ि' (short-i) is typed BEFORE its consonant cluster in Kruti Dev; move it
# to AFTER the cluster ((half-consonant)* + full consonant). Single pass.
_IRE = re.compile("ि((?:[%s]्)*[%s])" % (_CONS, _CONS))
# reph: a consonant followed by 'Z' means र् rides on that consonant cluster.
_ZRE = re.compile("([%s])([%s]*)Z" % (_CONS, _MATRAS))
# anusvara typed before a vowel matra must follow it (Unicode canonical order).
_NRE = re.compile("ं([ािीुूृेैोौ])")


def convert(s: str) -> str:
    """Convert a Kruti Dev 010 encoded string to Unicode Devanagari."""
    if not s:
        return s
    for a, b in PAIRS:
        s = s.replace(a, b)
    s = _ZRE.sub(lambda m: "र्" + m.group(1) + m.group(2), s)
    s = _IRE.sub(r"\1ि", s)
    s = _NRE.sub(r"\1ं", s)
    return s


def convert_docx(path: str) -> str:
    """Decode every paragraph and table cell of a .docx, in document order."""
    from docx import Document  # local import: only needed for the .docx path

    doc = Document(path)
    lines: list[str] = []
    for para in doc.paragraphs:
        lines.append(convert(para.text))
    for ti, table in enumerate(doc.tables):
        lines.append(f"\n[table {ti}: {len(table.rows)}x{len(table.columns)}]")
        for row in table.rows:
            cells = [convert(c.text).strip() for c in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def main(argv: list[str]) -> int:
    if argv and argv[0] == "--text":
        print(convert(" ".join(argv[1:])))
        return 0
    if argv:
        path = argv[0]
        if path.lower().endswith((".docx", ".doc")):
            print(convert_docx(path))
        else:
            with open(path, encoding="utf-8") as fh:
                print(convert(fh.read()))
        return 0
    # stdin
    print(convert(sys.stdin.read()))
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
