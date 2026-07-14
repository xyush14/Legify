"""Canonical draft → native Word .docx (high-fidelity, for the Word add-in).

The web render lays a court paper out with flexbox/grid CSS, which Word's HTML
import cannot reproduce — so converting that HTML into Word loses the court
format. For the Word add-in we instead author the SAME content as a NATIVE .docx
using Word's own constructs (centered court name, a stacked party header with
"बनाम", numbered "यह कि" grounds, right-aligned signature, real tables). The
add-in drops this file into the open document with `insertFileFromBase64`, so the
formatting lands 1:1 and stays fully editable.

Input is the canonical HTML fragment (`template_adapter.document`), parsed by its
known, fixed class vocabulary — .hdr-court, .hdr-party, .hdr-versus, .hdr-title,
.cb-prelude, .cb-paras, .cb-prayer, .cb-sig, .cb-table. One mapper serves every
canonical type because they all share this skeleton.

Devanagari: each run is given a complex-script (w:cs) font so Hindi renders in
Word even where the Latin font wouldn't cover it.

`docxtpl` path (future): when a real filed .docx template is supplied for a type,
we can fill it directly for pixel-perfect court format; this programmatic builder
is the always-available baseline.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH as ALIGN
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

import os

# Latin text uses a court-standard serif that exists on every OS. Devanagari is
# COMPLEX SCRIPT — Word renders it with the run's `w:cs` font, so that's where the
# Hindi font must go. Default "Nirmala UI" is the modern Windows Devanagari font
# (present on Win8+/Win11, unlike the legacy "Mangal", which recent Windows drops);
# on macOS Word substitutes an installed Devanagari face automatically. Override
# with HEADNOTE_DEVA_FONT if a deployment standardises on another font.
_LATIN_FONT = "Times New Roman"
_DEVA_FONT = os.environ.get("HEADNOTE_DEVA_FONT", "Nirmala UI")
_BASE_PT = 12


def _script_font(lang: str) -> str:
    # kept for the Normal-style default: Devanagari face for Hindi docs so any
    # stray un-run-styled text still resolves to a Devanagari-capable font.
    return _DEVA_FONT if lang != "en" else _LATIN_FONT


def _style_run(run, *, lang: str, bold=False, underline=False, size=_BASE_PT):
    """Style a run and pin fonts per script: Latin → Times New Roman (w:ascii/hAnsi),
    Devanagari → the Devanagari font (w:cs). Word chooses per character by Unicode
    range, so Hindi always lands on the complex-script font and renders."""
    run.bold = bold
    run.underline = underline
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), _LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), _LATIN_FONT)
    rfonts.set(qn("w:cs"), _DEVA_FONT)


_ALIGN = {"center": ALIGN.CENTER, "right": ALIGN.RIGHT, "justify": ALIGN.JUSTIFY, "left": ALIGN.LEFT}


def _add_runs(para, node, *, lang, bold=False, underline=False, size=_BASE_PT):
    """Walk an HTML node's children into runs on `para`, honouring <br>, <b>/<strong>,
    <u>, and .ph placeholder spans (rendered underlined so blanks stay visible)."""
    from bs4 import NavigableString, Tag
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text.strip() == "" and "\n" in text:
                continue
            if text:
                _style_run(para.add_run(text), lang=lang, bold=bold, underline=underline, size=size)
        elif isinstance(child, Tag):
            name = child.name.lower()
            if name == "br":
                para.add_run().add_break()
            elif name in ("b", "strong"):
                _add_runs(para, child, lang=lang, bold=True, underline=underline, size=size)
            elif name in ("u",):
                _add_runs(para, child, lang=lang, bold=bold, underline=True, size=size)
            elif name == "span" and "ph" in (child.get("class") or []):
                # unfilled placeholder → keep the hint text, underlined, so it's obvious
                _style_run(para.add_run(child.get_text()), lang=lang, bold=bold, underline=True, size=size)
            else:
                _add_runs(para, child, lang=lang, bold=bold, underline=underline, size=size)


def _para(doc, node, *, lang, align="left", bold=False, underline=False, size=_BASE_PT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = _ALIGN.get(align, ALIGN.LEFT)
    p.paragraph_format.space_after = Pt(space_after)
    if node is None:
        return p
    if isinstance(node, str):
        _style_run(p.add_run(node), lang=lang, bold=bold, underline=underline, size=size)
    else:
        _add_runs(p, node, lang=lang, bold=bold, underline=underline, size=size)
    return p


def _emit_table(doc, tbl, *, lang):
    rows = tbl.find_all("tr")
    if not rows:
        return
    ncol = max(len(r.find_all(["td", "th"])) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    table.style = "Table Grid"
    for ri, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"])
        for ci in range(ncol):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].text = ""  # clear default empty run
            if ci < len(cells):
                _add_runs(cell.paragraphs[0], cells[ci], lang=lang,
                          bold=(cells[ci].name == "th"), size=_BASE_PT - 1)


def _emit_grounds(doc, ol, *, lang):
    """Numbered "यह कि …" grounds — one justified, hanging-indent paragraph each."""
    for i, li in enumerate(ol.find_all("li", recursive=False), start=1):
        p = doc.add_paragraph()
        p.alignment = ALIGN.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        _style_run(p.add_run(f"{i}. "), lang=lang, size=_BASE_PT)
        _add_runs(p, li, lang=lang, size=_BASE_PT)


def _dispatch_block(doc, el, cls, *, lang) -> bool:
    """Emit `el` if it's a recognized leaf block; return True if handled (caller
    then does NOT recurse into it), False if it's a container to descend into."""
    if "hdr-side" in cls:
        _para(doc, el, lang=lang, align="right", size=10, space_after=2); return True
    if "hdr-court" in cls:
        _para(doc, el, lang=lang, align="center", bold=True, size=13, space_after=4); return True
    if "hdr-case" in cls:
        _para(doc, el, lang=lang, align="center", size=11, space_after=8); return True
    if "hdr-party-label" in cls:
        _para(doc, el, lang=lang, align="left", bold=True, size=11, space_after=0); return True
    if "hdr-party-desc" in cls:
        _para(doc, el, lang=lang, align="left", size=11, space_after=6); return True
    if "hdr-versus" in cls:
        _para(doc, el, lang=lang, align="center", bold=True, size=11, space_after=6); return True
    if "hdr-title" in cls:
        _para(doc, el, lang=lang, align="center", bold=True, underline=True, size=12, space_after=10); return True
    if "cb-prelude" in cls:
        _para(doc, el, lang=lang, align="justify", space_after=8); return True
    if "cb-prayer" in cls:
        _para(doc, el, lang=lang, align="justify", space_after=8); return True
    if "cb-block-label" in cls:
        _para(doc, el, lang=lang, align="left", bold=True, space_after=2); return True
    if ("l" in cls or "r" in cls) and el.find_parent(class_="cb-sig") is not None:
        _para(doc, el, lang=lang, align=("right" if "r" in cls else "left"), size=11, space_after=2); return True
    return False


def _walk(doc, el, *, lang):
    """Recursively emit `el` in document order. Recognized leaf blocks / ol / table
    are emitted whole and not descended into; everything else is a container we
    recurse through, so paragraph order matches the source exactly."""
    from bs4 import Tag
    name = (el.name or "").lower()
    cls = el.get("class") or []
    if name == "ol" and "cb-paras" in cls:
        _emit_grounds(doc, el, lang=lang); return
    if name == "table" and "cb-table" in cls:
        _emit_table(doc, el, lang=lang); return
    if _dispatch_block(doc, el, cls, lang=lang):
        return
    for child in el.children:
        if isinstance(child, Tag):
            _walk(doc, child, lang=lang)


def html_fragment_to_docx(body_html: str, lang: str = "hi") -> bytes:
    """Canonical body HTML → native .docx bytes."""
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(body_html or "", "html.parser")
    sheet = soup.select_one(".doc-a4") or soup

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.25)
    sec.right_margin = Inches(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = _script_font(lang)
    normal.font.size = Pt(_BASE_PT)

    _walk(doc, sheet, lang=lang)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def canonical_to_docx(doc_type: str, fields: dict, lang: str = "hi") -> bytes:
    """Render a canonical draft directly to native .docx bytes."""
    from headnote.drafter import template_adapter as TA
    html = TA.document(doc_type, fields or {}, lang)
    body = html.split("</style>", 1)[-1]
    return html_fragment_to_docx(body, lang)
